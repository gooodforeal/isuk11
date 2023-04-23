import pandas as pd
from datetime import datetime
from docx import Document
import re
import os
from barcode.writer import ImageWriter
from barcode.codex import Code128
from docx2pdf import convert


def replace_text(paragraph, key, value):
    "Переставляет текстовые переменные в шаблоне"
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


def read_data():
    "Читает данные из файла excel в dict"
    data = {}
    df = pd.read_excel("data.xlsx")
    df = df.sort_values("Грузополучатель")
    for i, row in df.iterrows():
        data[row["Грузополучатель"]] = {}
        data[row["Грузополучатель"]]["Позиции"] = []
        data[row["Грузополучатель"]]["{{NUMBER}}"] = len(list(data.keys()))
    for i, row in df.iterrows():
        data[row["Грузополучатель"]]["{{TO}}"] = row["Грузополучатель"]
        data[row["Грузополучатель"]]["{{FROM}}"] = row["Грузоотправитель"]
        data[row["Грузополучатель"]]["{{DOCUMENT}}"] = row["Основание"]
        data[row["Грузополучатель"]]["{{DATE}}"] = datetime.today().date().strftime("%d.%m.%y")
        data[row["Грузополучатель"]]["Позиции"].append([len(data[row["Грузополучатель"]]["Позиции"]) + 1, row["Товар"], row["Ед."],  row["Кол-во"]])
        data[row["Грузополучатель"]]["{{DOCUMENT}}"] = row["Основание"]
    return data


def main():
    "Основная функция"
    data = read_data()
    pattern = r"{{\w{1,}}}"
    template = 'template.docx'
    n = 1
    for key, value in data.items():
        result = f'накладная_{n:02}_{key}'
        template_doc = Document(template)
        barcode_1 = Code128(f'AB{data[key]["{{NUMBER}}"]}_{data[key]["{{DATE}}"]}', writer=ImageWriter(format='PNG'))
        filename = barcode_1.save('example', options=dict(module_height=3))
        par = template_doc.paragraphs[0].insert_paragraph_before()
        lr = par.add_run()
        lr.add_picture('example.PNG')
        os.remove('example.PNG')
        for k, v in value.items():
            if re.fullmatch(pattern, k.strip()):
                for paragraph in template_doc.paragraphs:
                    replace_text(paragraph, k, str(v))
        my_table = template_doc.tables[0]
        for line in value["Позиции"]:
            row = my_table.add_row()
            i = 0
            for cell in row.cells:
                cell.text = str(line[i])
                i += 1
        template_doc.save(f"result/{result}.docx")
        convert(f"result/{result}.docx", f"result/{result}.pdf")
        os.remove(f"result/{result}.docx")
        n += 1


if __name__ == "__main__":
    main()
