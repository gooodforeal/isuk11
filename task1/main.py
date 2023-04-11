from docx import Document
import re
from datetime import datetime


def read_data():
    data = {}
    pattern = r"{{\w{1,}}}"
    keys = []
    values = []
    durations = []
    with open("data.txt", "r", encoding="UTF-8") as f:
        lines = f.readlines()
        split_index = lines.index("\n")
        items = lines[:split_index]
        dates = lines[split_index + 1:]
        for line in items:
            if re.fullmatch(pattern, line.strip()):
                keys.append(line.strip())
            else:
                values.append(line.strip())
        for date in dates:
            date1, date2 = date.split(" ")
            date1 = date1.strip()
            date2 = date2.strip()
            dif = str((datetime.strptime(date2, "%d.%m.%Y") - datetime.strptime(date1, "%d.%m.%Y")).days)
            durations.append([date1, date2, dif])
        for i in range(len(keys)):
            if keys[i] == "{{DATE}}":
                data[keys[i]] = datetime.today().date().strftime("%d.%m.%y")
            else:
                data[keys[i]] = values[i]
        data["DURATIONS"] = durations
    return data


def main():
    data = read_data()
    template = 'template.docx'
    result = 'result.docx'
    template_doc = Document(template)
    for key, value in data.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, key, value)
    my_table = template_doc.tables[0]
    n = 1
    for line in data["DURATIONS"]:
        row = my_table.add_row()
        i = -1
        for cell in row.cells:
            if i == -1:
                cell.text = str(n)
            else:
                cell.text = line[i]
            i += 1
        n += 1
    template_doc.save(result)


def replace_text(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


main()
